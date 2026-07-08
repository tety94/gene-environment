#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generate_table1.py
===================

Generates paper "Table 1" (Cohort 1 vs Cohort 2 comparison) from:

  - componenti_ambientali_1.csv   -> cohort 1 data (demographics + environmental variables)
  - componenti_ambientali_2.csv   -> cohort 2 data (demographics + environmental variables)
  - componenti_ambientali_fumo.csv        -> fumo_boolean (ids shared by cohort 1 and 2)
  - componenti_ambientali_scolarita.csv   -> education_level / education_years (shared ids)

Logic:
  1. componenti_ambientali_1.csv and _2.csv are concatenated and de-duplicated
     on "id" (rows with an implausible id, e.g. free-text placeholders like
     "novara" or "not collected", are discarded BEFORE de-duplication so they
     don't collapse different patients into a single row).
  2. Real cohort membership (Cohort 1 / Cohort 2) is derived from which of the
     two genotype files an id (doubled: "id_id") appears in, NOT from any
     column inside the environmental files. If an id appears in both genotype
     files, Cohort 1 wins.
  3. Each cohort is enriched (LEFT JOIN on "id") with the *new* columns found
     in the fumo/scolarita files (duplicated demographic columns such as
     sex/onset_site/onset_age/diagnostic_delay are not re-imported, but are
     checked for consistency and logged if discordant).
  4. For each variable:
       - CONTINUOUS  -> median [IQR], Mann-Whitney U test
       - CATEGORICAL -> n (%), Chi-square test (or Fisher's exact if expected <5)
  5. Output in the "output/" folder:
       - table1.csv
       - table1.docx      (ready-to-use table for the paper, with figures in appendix)
       - images/*.png     (boxplots, bar charts, buffer-distance p-value trend)
       - run_log.txt      (join/merge log, missing data, warnings)

Usage:
    python generate_table1.py --input-dir . --output-dir output

    # Use all 6 buffer distances for "seminativi" instead of the 1500m default:
    python generate_table1.py --input-dir . --output-dir output --seminativi-buffers all

    # Use a custom subset of buffer distances for "seminativi":
    python generate_table1.py --input-dir . --output-dir output --seminativi-buffers 500,1500

Dependencies:
    pip install pandas scipy matplotlib python-docx --break-system-packages
"""

import argparse
import os
import re
import sys
import warnings
from datetime import datetime

import numpy as np
import pandas as pd
from scipy.stats import mannwhitneyu, chi2_contingency, fisher_exact

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

warnings.filterwarnings("ignore")

# ----------------------------------------------------------------------------
# CONFIG
# ----------------------------------------------------------------------------

COHORT1_FILE = "componenti_ambientali_1.csv"
COHORT2_FILE = "componenti_ambientali_2.csv"
FUMO_FILE = "componenti_ambientali_fumo.csv"
SCOLARITA_FILE = "componenti_ambientali_scolarita.csv"

ID_COL = "id"

# Genotype files that define REAL cohort membership (the id in
# componenti_ambientali_1/2.csv matches EXACTLY, doubled, the id in these
# files: e.g. "RES02977_RES02977"). Cohort 1 = ids found in the gen1 file
# (RES...), Cohort 2 = ids found in the gen2 file (ACH...). If an id is found
# in both, Cohort 1 wins.
GEN1_IDS_FILE_DEFAULT = "/mnt/cresla_prod/genome_datasets/merged_csv/full_chr_gen1_test1.csv"
GEN2_IDS_FILE_DEFAULT = "/srv/python-projects/gene-environment/output_combined_risaie.csv"

# Computed column used for the Table 1 comparison ("Cohort 1"/"Cohort 2"),
# derived from genotype-file membership above, NOT from parals_codals (which
# is just an arbitrary/id-like label and is treated as a regular descriptive
# categorical variable).
GROUP_COL = "cohort"
GROUP1_LABEL = "Cohort 1"
GROUP2_LABEL = "Cohort 2"

# Continuous demographic/clinical variables expected in the cohort files (if present)
BASE_CONTINUOUS_VARS = ["onset_age", "diagnostic_delay", "survival"]

# Categorical demographic/clinical variables expected in the cohort files (if present)
BASE_CATEGORICAL_VARS = ["sex", "onset_site", "parals_codals"]

# "Demographic" columns duplicated in the fumo/scolarita files: NOT imported
# during the join (the cohort-file version is kept), but compared for
# consistency and logged if discordant.
DEMOGRAPHIC_OVERLAP_COLS = ["sex", "onset_site", "onset_age", "diagnostic_delay"]

# Pattern to auto-detect environmental variables (6 buffers x 3 categories)
BUFFER_VAR_PATTERN = re.compile(r"^(seminativi|vigneti|risaie)_(\d+)$")

# Some categorical columns (e.g. parals_codals) can contain leftover
# "code" values (e.g. "COD-0001", "PAR-2008-056") that are not real
# categories but ids/codes that ended up in the column by mistake. These
# values are treated as missing (NaN) for that specific variable only,
# without excluding the patient from the analysis of other variables.
CODE_LIKE_PATTERN = re.compile(r"^[A-Za-z]{2,10}(-\d{2,6}){1,3}$")
COLUMNS_TO_CLEAN_CODE_LIKE = ["parals_codals"]

# Default buffer distances (meters) used for "seminativi" variables. Unlike
# vigneti/risaie (which always use all 6 buffers), seminativi defaults to a
# single representative distance to keep Table 1 and the figures focused.
# Override with --seminativi-buffers (comma-separated list, or "all").
DEFAULT_SEMINATIVI_BUFFERS = [1500]

ALPHA = 0.05


def load_ids_from_genotype_file(path, log, label):
    """Reads ONLY the first field (id) of every line in a genotype file,
    without parsing the thousands of following SNP columns. Efficient even
    on very large files because it stops at the first comma of each line."""
    if not os.path.exists(path):
        log.add(f"ERROR: genotype file not found for {label} -> {path}")
        return set()
    ids = set()
    n_lines = 0
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        f.readline()  # header
        for line in f:
            n_lines += 1
            comma_idx = line.find(",")
            id_val = line[:comma_idx].strip() if comma_idx != -1 else line.strip()
            if id_val:
                ids.add(id_val)
    log.add(f"Genotype file {label} ({path}): {n_lines} rows, {len(ids)} unique ids")
    return ids


# ----------------------------------------------------------------------------
# UTILITY
# ----------------------------------------------------------------------------

class RunLog:
    """Collects messages to be written later to run_log.txt"""

    def __init__(self):
        self.lines = []

    def add(self, msg):
        stamp = datetime.now().strftime("%H:%M:%S")
        line = f"[{stamp}] {msg}"
        print(line)
        self.lines.append(line)

    def save(self, path):
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(self.lines))


def read_csv_robust(path, log):
    """Reads a csv handling quoting automatically and preparing numeric
    columns for downstream conversion."""
    if not os.path.exists(path):
        log.add(f"ERROR: file not found -> {path}")
        return None
    df = pd.read_csv(path, sep=",", quotechar='"', skipinitialspace=True)
    df.columns = [c.strip().strip('"') for c in df.columns]
    if ID_COL not in df.columns:
        log.add(f"WARNING: column '{ID_COL}' not found in {path} (columns: {list(df.columns)})")
    else:
        df[ID_COL] = df[ID_COL].astype(str).str.strip()
    log.add(f"Read {path}: {df.shape[0]} rows, {df.shape[1]} columns")
    return df


def parse_buffer_list(value):
    """Parses --seminativi-buffers: 'all' -> None (no restriction), or a
    comma-separated list of ints, e.g. '500,1500' -> [500, 1500]."""
    if value is None:
        return None
    if value.strip().lower() == "all":
        return None
    return sorted(int(x.strip()) for x in value.split(",") if x.strip())


def detect_buffer_vars(df, seminativi_buffers=None):
    """Returns the list of environmental columns seminativi/vigneti/risaie_<buffer>
    found in the dataframe, sorted by category and buffer distance.
    seminativi_buffers: None = keep all seminativi buffers; otherwise a list
    of allowed buffer distances (ints) restricting seminativi_* only
    (vigneti/risaie are never restricted)."""
    found = [c for c in df.columns if BUFFER_VAR_PATTERN.match(c)]

    if seminativi_buffers is not None:
        allowed = set(seminativi_buffers)
        found = [
            c for c in found
            if not c.startswith("seminativi_") or int(c.split("_")[1]) in allowed
        ]

    def sort_key(c):
        m = BUFFER_VAR_PATTERN.match(c)
        return (m.group(1), int(m.group(2)))

    return sorted(found, key=sort_key)


def check_demographic_consistency(cohort_df, side_df, side_name, cohort_name, log):
    """Compares shared demographic columns (matched by id) between the cohort
    file and the fumo/scolarita file, and logs any discrepancies."""
    common_cols = [c for c in DEMOGRAPHIC_OVERLAP_COLS if c in cohort_df.columns and c in side_df.columns]
    if not common_cols:
        return
    merged = cohort_df[[ID_COL] + common_cols].merge(
        side_df[[ID_COL] + common_cols], on=ID_COL, how="inner", suffixes=("_cohort", f"_{side_name}")
    )
    n_mismatch_total = 0
    for col in common_cols:
        c1 = merged[f"{col}_cohort"]
        c2 = merged[f"{col}_{side_name}"]
        mismatch = merged[(c1.astype(str) != c2.astype(str)) & c1.notna() & c2.notna()]
        if len(mismatch) > 0:
            n_mismatch_total += len(mismatch)
            log.add(
                f"WARNING [{cohort_name} vs {side_name}]: {len(mismatch)} ids with discordant "
                f"values for '{col}' (cohort-file value kept)."
            )
    if n_mismatch_total == 0:
        log.add(f"OK: no demographic discrepancy between {cohort_name} and {side_name} on columns {common_cols}")


def merge_side_file(cohort_df, side_df, side_name, cohort_name, log):
    """Left-join of cohort_df with only the NEW columns present in side_df
    (everything except id and demographic columns already in cohort_df)."""
    if side_df is None:
        return cohort_df

    check_demographic_consistency(cohort_df, side_df, side_name, cohort_name, log)

    new_cols = [c for c in side_df.columns if c not in cohort_df.columns and c != ID_COL]
    if not new_cols:
        log.add(f"WARNING: no new column to import from {side_name} into {cohort_name}")
        return cohort_df

    slim = side_df[[ID_COL] + new_cols].drop_duplicates(subset=ID_COL)
    before = cohort_df.shape[0]
    merged = cohort_df.merge(slim, on=ID_COL, how="left")
    n_missing = merged[new_cols[0]].isna().sum()
    log.add(
        f"Join {cohort_name} + {side_name}: {before} rows -> {merged.shape[0]} rows; "
        f"columns added: {new_cols}; ids with no match in {side_name}: {n_missing}"
    )
    return merged


def clean_code_like_values(df, columns, log):
    """Replaces with NaN values that look like codes/ids (e.g. 'COD-0001',
    'PAR-2008-056') instead of real categories, in the given columns."""
    for col in columns:
        if col not in df.columns:
            continue
        s = df[col].astype(str)
        mask = s.str.match(CODE_LIKE_PATTERN, na=False)
        n_bad = mask.sum()
        if n_bad > 0:
            examples = df.loc[mask, col].astype(str).unique().tolist()[:5]
            remaining = df.loc[~mask, col].dropna().astype(str).unique().tolist()
            log.add(f"Cleaning '{col}': {n_bad} code-like values treated as missing "
                     f"(examples: {examples}); remaining valid categories: {remaining}")
            df.loc[mask, col] = np.nan
    return df


def is_plausible_patient_id(id_val):
    """A plausible patient id has no whitespace, is not empty/nan, and
    contains at least one digit (placeholders like 'novara', 'not collected',
    'novara not collected' are words with no digits, often with spaces, and
    are discarded)."""
    if id_val is None:
        return False
    s = str(id_val).strip()
    if not s or s.lower() in ("nan", "none", "null", ""):
        return False
    if any(ch.isspace() for ch in s):
        return False
    if not any(ch.isdigit() for ch in s):
        return False
    return True


# ----------------------------------------------------------------------------
# STATISTICS
# ----------------------------------------------------------------------------

def fmt_p(p):
    if p is None or (isinstance(p, float) and np.isnan(p)):
        return "n/a"
    if p < 0.001:
        return "<0.001"
    return f"{p:.3f}"


def continuous_summary(series):
    s = pd.to_numeric(series, errors="coerce").dropna()
    n = len(s)
    if n == 0:
        return "n=0", n
    med, q1, q3 = s.median(), s.quantile(0.25), s.quantile(0.75)
    return f"{med:.1f} [{q1:.1f}\u2013{q3:.1f}]", n


def mannwhitney_p(s1, s2):
    s1 = pd.to_numeric(s1, errors="coerce").dropna()
    s2 = pd.to_numeric(s2, errors="coerce").dropna()
    if len(s1) < 1 or len(s2) < 1:
        return np.nan
    try:
        _, p = mannwhitneyu(s1, s2, alternative="two-sided")
        return p
    except ValueError:
        return np.nan


def categorical_counts(series):
    s = series.dropna()
    counts = s.value_counts()
    total = counts.sum()
    return counts, total


def chi2_or_fisher_p(df, var, group_col):
    ct = pd.crosstab(df[var], df[group_col])
    ct = ct.loc[:, (ct.sum(axis=0) > 0)]
    if ct.shape[0] < 2 or ct.shape[1] < 2:
        return "n/a", np.nan
    try:
        chi2, p, dof, expected = chi2_contingency(ct, correction=False)
    except ValueError:
        return "n/a", np.nan
    if ct.shape == (2, 2):
        if (expected < 5).any():
            try:
                _, p = fisher_exact(ct.values)
                return "Fisher's exact", p
            except ValueError:
                return "n/a", np.nan
        return "Chi-square", p
    test_name = "Chi-square"
    if (expected < 5).sum() > 0:
        test_name = "Chi-square (warning: some expected cells <5)"
    return test_name, p


# ----------------------------------------------------------------------------
# TABLE 1 CONSTRUCTION
# ----------------------------------------------------------------------------

def build_table1(df, group_col, continuous_vars, categorical_vars, log):
    """Returns a list of dicts, one per table row (variable header + any
    category sub-rows)."""
    groups = sorted(df[group_col].dropna().unique())
    if len(groups) != 2:
        log.add(f"WARNING: expected 2 groups in '{group_col}', found: {groups}")
    g1, g2 = groups[0], groups[1]

    rows = []
    n1_tot = (df[group_col] == g1).sum()
    n2_tot = (df[group_col] == g2).sum()
    rows.append({
        "Variable": "Total N, n",
        "Category": "",
        f"{g1}": f"{n1_tot}",
        f"{g2}": f"{n2_tot}",
        "p-value": "",
        "Test": "",
        "significant": False,
    })

    for var in continuous_vars:
        if var not in df.columns:
            log.add(f"WARNING: continuous variable '{var}' not present, skipped")
            continue
        s1 = df.loc[df[group_col] == g1, var]
        s2 = df.loc[df[group_col] == g2, var]
        summ1, n1 = continuous_summary(s1)
        summ2, n2 = continuous_summary(s2)
        p = mannwhitney_p(s1, s2)
        rows.append({
            "Variable": f"{var}, median [IQR]",
            "Category": "",
            f"{g1}": f"{summ1} (n={n1})",
            f"{g2}": f"{summ2} (n={n2})",
            "p-value": fmt_p(p),
            "Test": "Mann-Whitney U",
            "significant": (p is not None and not np.isnan(p) and p < ALPHA),
        })

    for var in categorical_vars:
        if var not in df.columns:
            log.add(f"WARNING: categorical variable '{var}' not present, skipped")
            continue
        c1, tot1 = categorical_counts(df.loc[df[group_col] == g1, var])
        c2, tot2 = categorical_counts(df.loc[df[group_col] == g2, var])
        test_name, p = chi2_or_fisher_p(df, var, group_col)
        sig = (p is not None and not np.isnan(p) and p < ALPHA)
        rows.append({
            "Variable": f"{var}, n (%)",
            "Category": "",
            f"{g1}": f"n={tot1}",
            f"{g2}": f"n={tot2}",
            "p-value": fmt_p(p),
            "Test": test_name,
            "significant": sig,
        })
        categories = sorted(set(c1.index.tolist()) | set(c2.index.tolist()), key=lambda x: str(x))
        for cat in categories:
            n1 = c1.get(cat, 0)
            n2 = c2.get(cat, 0)
            pct1 = (n1 / tot1 * 100) if tot1 else 0
            pct2 = (n2 / tot2 * 100) if tot2 else 0
            rows.append({
                "Variable": "",
                "Category": f"   {cat}",
                f"{g1}": f"{n1} ({pct1:.1f}%)",
                f"{g2}": f"{n2} ({pct2:.1f}%)",
                "p-value": "",
                "Test": "",
                "significant": False,
            })

    return rows, g1, g2


# ----------------------------------------------------------------------------
# CSV EXPORT
# ----------------------------------------------------------------------------

def save_table1_csv(rows, g1, g2, path):
    df_out = pd.DataFrame(rows)
    cols = ["Variable", "Category", str(g1), str(g2), "p-value", "Test"]
    cols = [c for c in cols if c in df_out.columns]
    df_out = df_out[cols]
    df_out.to_csv(path, index=False, encoding="utf-8-sig")


# ----------------------------------------------------------------------------
# DOCX EXPORT
# ----------------------------------------------------------------------------

def _set_cell_shading(cell, color_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def save_table1_docx(rows, g1, g2, path, images_for_appendix=None):
    doc = Document()

    title = doc.add_heading("Table 1. Demographic, clinical and environmental characteristics "
                             "by study cohort", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitle = doc.add_paragraph()
    run = subtitle.add_run(
        f"Comparison {g1} vs {g2}. Continuous variables: median [IQR], Mann-Whitney U test. "
        f"Categorical variables: n (%), Chi-square or Fisher's exact test (expected cells <5). "
        f"p < {ALPHA} highlighted in bold."
    )
    run.italic = True
    run.font.size = Pt(9)

    headers = ["Variable", str(g1), str(g2), "p-value", "Test"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
        _set_cell_shading(hdr_cells[i], "D9E2F3")

    for row in rows:
        cells = table.add_row().cells
        var_label = row["Variable"] if row["Variable"] else row["Category"]
        values = [var_label, row.get(str(g1), ""), row.get(str(g2), ""), row["p-value"], row["Test"]]
        for i, v in enumerate(values):
            cells[i].text = str(v)
            if row.get("significant") and i == 3:
                for p in cells[i].paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            if not row["Variable"] and row["Category"] and i == 0:
                for p in cells[i].paragraphs:
                    for r in p.runs:
                        r.italic = True

    for col_idx, width in enumerate([Inches(2.3), Inches(1.5), Inches(1.5), Inches(0.9), Inches(1.6)]):
        for row in table.rows:
            row.cells[col_idx].width = width

    if images_for_appendix:
        doc.add_page_break()
        doc.add_heading("Appendix - Figures", level=1)
        for img_path, caption in images_for_appendix:
            if os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(6))
                cap = doc.add_paragraph(caption)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in cap.runs:
                    r.italic = True
                    r.font.size = Pt(9)

    doc.save(path)


# ----------------------------------------------------------------------------
# FIGURES
# ----------------------------------------------------------------------------

def plot_continuous_box(df, var, group_col, g1, g2, out_path, log):
    s1 = pd.to_numeric(df.loc[df[group_col] == g1, var], errors="coerce").dropna()
    s2 = pd.to_numeric(df.loc[df[group_col] == g2, var], errors="coerce").dropna()
    if len(s1) == 0 or len(s2) == 0:
        return None
    p = mannwhitney_p(s1, s2)
    fig, ax = plt.subplots(figsize=(5, 4.5))
    try:
        bp = ax.boxplot([s1, s2], tick_labels=[str(g1), str(g2)], patch_artist=True, showmeans=True)
    except TypeError:
        bp = ax.boxplot([s1, s2], labels=[str(g1), str(g2)], patch_artist=True, showmeans=True)
    colors = ["#4C72B0", "#DD8452"]
    for patch, c in zip(bp["boxes"], colors):
        patch.set_facecolor(c)
        patch.set_alpha(0.6)
    ax.set_title(f"{var}\nMann-Whitney U p={fmt_p(p)}")
    ax.set_ylabel(var)
    fig.tight_layout()
    fig.savefig(out_path, dpi=150)
    plt.close(fig)
    return out_path


def plot_buffer_grouped_box(df, category, group_col, g1, g2, out_dir, log, buffer_cols=None):
    """category in {'seminativi','vigneti','risaie'}: boxplot grouped by
    buffer distance, side-by-side cohort1/cohort2. If buffer_cols is given,
    only those columns are plotted (used to respect --seminativi-buffers);
    otherwise all matching columns in df are auto-detected."""
    if buffer_cols is None:
        buffer_cols = sorted(
            [c for c in df.columns if c.startswith(category + "_")],
            key=lambda c: int(c.split("_")[1]),
        )
    else:
        buffer_cols = sorted(buffer_cols, key=lambda c: int(c.split("_")[1]))

    if not buffer_cols:
        return None
    distances = [c.split("_")[1] for c in buffer_cols]

    fig, ax = plt.subplots(figsize=(max(6, len(buffer_cols) * 1.1), 5))
    positions1, positions2 = [], []
    data1, data2 = [], []
    width = 0.35
    for i, col in enumerate(buffer_cols):
        s1 = pd.to_numeric(df.loc[df[group_col] == g1, col], errors="coerce").dropna()
        s2 = pd.to_numeric(df.loc[df[group_col] == g2, col], errors="coerce").dropna()
        data1.append(s1)
        data2.append(s2)
        positions1.append(i - width / 2)
        positions2.append(i + width / 2)

    bp1 = ax.boxplot(data1, positions=positions1, widths=width, patch_artist=True, showfliers=False)
    bp2 = ax.boxplot(data2, positions=positions2, widths=width, patch_artist=True, showfliers=False)
    for patch in bp1["boxes"]:
        patch.set_facecolor("#4C72B0")
        patch.set_alpha(0.6)
    for patch in bp2["boxes"]:
        patch.set_facecolor("#DD8452")
        patch.set_alpha(0.6)

    ax.set_xticks(range(len(buffer_cols)))
    ax.set_xticklabels([f"{d}m" for d in distances])
    ax.set_xlabel("Buffer (meters)")
    ax.set_ylabel(f"{category} (value)")
    ax.set_title(f"{category}: {g1} vs {g2} comparison by buffer")
    ax.legend([bp1["boxes"][0], bp2["boxes"][0]], [str(g1), str(g2)], loc="best")
    fig.tight_layout()
    out_path = os.path.join(out_dir, f"boxplot_{category}_per_buffer.png")
    fig.savefig(out_path, dpi=150)
    plt.close(fig)
    return out_path


def plot_buffer_pvalue_trend(df, category, group_col, g1, g2, out_dir, log, buffer_cols=None):
    if buffer_cols is None:
        buffer_cols = sorted(
            [c for c in df.columns if c.startswith(category + "_")],
            key=lambda c: int(c.split("_")[1]),
        )
    else:
        buffer_cols = sorted(buffer_cols, key=lambda c: int(c.split("_")[1]))

    if not buffer_cols:
        return None
    distances, pvals = [], []
    for col in buffer_cols:
        s1 = df.loc[df[group_col] == g1, col]
        s2 = df.loc[df[group_col] == g2, col]
        p = mannwhitney_p(s1, s2)
        distances.append(int(col.split("_")[1]))
        pvals.append(p)

    fig, ax = plt.subplots(figsize=(6, 4.5))
    ax.plot(distances, pvals, marker="o", color="#4C72B0")
    ax.axhline(ALPHA, color="red", linestyle="--", linewidth=1, label=f"threshold p={ALPHA}")
    ax.set_xlabel("Buffer (meters)")
    ax.set_ylabel("p-value (Mann-Whitney U)")
    ax.set_title(f"{category}: p-value trend across buffer distances")
    ax.legend()
    fig.tight_layout()
    out_path = os.path.join(out_dir, f"pvalue_trend_{category}.png")
    fig.savefig(out_path, dpi=150)
    plt.close(fig)
    return out_path


def plot_categorical_bar(df, var, group_col, g1, g2, out_dir, log):
    if var not in df.columns:
        return None
    c1, tot1 = categorical_counts(df.loc[df[group_col] == g1, var])
    c2, tot2 = categorical_counts(df.loc[df[group_col] == g2, var])
    categories = sorted(set(c1.index.tolist()) | set(c2.index.tolist()), key=lambda x: str(x))
    if not categories:
        return None
    pct1 = [c1.get(cat, 0) / tot1 * 100 if tot1 else 0 for cat in categories]
    pct2 = [c2.get(cat, 0) / tot2 * 100 if tot2 else 0 for cat in categories]

    x = np.arange(len(categories))
    width = 0.35
    fig, ax = plt.subplots(figsize=(max(5, len(categories) * 1.2), 4.5))
    ax.bar(x - width / 2, pct1, width, label=str(g1), color="#4C72B0", alpha=0.8)
    ax.bar(x + width / 2, pct2, width, label=str(g2), color="#DD8452", alpha=0.8)
    ax.set_xticks(x)
    ax.set_xticklabels([str(c) for c in categories], rotation=0)
    ax.set_ylabel("% within cohort")
    ax.set_title(f"{var}: distribution by cohort")
    ax.legend()
    fig.tight_layout()
    out_path = os.path.join(out_dir, f"barplot_{var}.png")
    fig.savefig(out_path, dpi=150)
    plt.close(fig)
    return out_path


# ----------------------------------------------------------------------------
# MAIN
# ----------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Generate Table 1 (Cohort 1 vs Cohort 2) with statistical tests and figures")
    parser.add_argument("--input-dir", default=".", help="Folder containing the 4 environmental/demographic input csv files")
    parser.add_argument("--output-dir", default="output", help="Output folder")
    parser.add_argument("--gen1-ids-file", default=GEN1_IDS_FILE_DEFAULT,
                         help="Genotype file whose ids define Cohort 1 (RES...)")
    parser.add_argument("--gen2-ids-file", default=GEN2_IDS_FILE_DEFAULT,
                         help="Genotype file whose ids define Cohort 2 (ACH...)")
    parser.add_argument("--seminativi-buffers", default="1500",
                         help="Comma-separated buffer distances (meters) to use for 'seminativi' "
                              "variables, or 'all' for all 6 buffers. Default: 1500. "
                              "vigneti/risaie always use all 6 buffers.")
    args = parser.parse_args()

    out_dir = args.output_dir
    img_dir = os.path.join(out_dir, "images")
    os.makedirs(out_dir, exist_ok=True)
    os.makedirs(img_dir, exist_ok=True)

    seminativi_buffers = parse_buffer_list(args.seminativi_buffers)

    log = RunLog()
    log.add("Starting Table 1 generation")
    if seminativi_buffers is None:
        log.add("seminativi buffers: ALL (6 distances)")
    else:
        log.add(f"seminativi buffers: {seminativi_buffers} (use --seminativi-buffers all for all 6)")

    cohort1 = read_csv_robust(os.path.join(args.input_dir, COHORT1_FILE), log)
    cohort2 = read_csv_robust(os.path.join(args.input_dir, COHORT2_FILE), log)
    fumo = read_csv_robust(os.path.join(args.input_dir, FUMO_FILE), log)
    scolarita = read_csv_robust(os.path.join(args.input_dir, SCOLARITA_FILE), log)

    if cohort1 is None or cohort2 is None:
        log.add("FATAL ERROR: cohort files missing, cannot proceed")
        log.save(os.path.join(out_dir, "run_log.txt"))
        sys.exit(1)

    # componenti_ambientali_1 and _2 are merged (they contain the same set of
    # patients/columns); if an id appears in both, only one row is kept.
    raw_combined = pd.concat([cohort1, cohort2], ignore_index=True, sort=False)

    # Some rows have a NON-plausible id (placeholder such as "novara",
    # "not collected", typically for patients with no genetic sample drawn).
    # If we don't discard them BEFORE de-duplication, different patients
    # sharing the same placeholder as "id" would collapse into a single row,
    # losing real data. Since they don't have a valid id anyway, they cannot
    # be matched to any genotype cohort, so they are explicitly excluded here
    # and logged.
    valid_mask = raw_combined[ID_COL].apply(is_plausible_patient_id)
    n_invalid = (~valid_mask).sum()
    if n_invalid > 0:
        invalid_examples = raw_combined.loc[~valid_mask, ID_COL].astype(str).unique().tolist()[:10]
        log.add(f"WARNING: {n_invalid} rows with a NON-plausible id (placeholder, e.g. {invalid_examples}) "
                 f"excluded BEFORE de-duplication so they don't collapse together; these patients don't "
                 f"have a valid id to match against a genotype file anyway.")
        raw_combined = raw_combined.loc[valid_mask].copy()

    n_before = raw_combined.shape[0]
    dup_mask = raw_combined.duplicated(subset=ID_COL, keep=False)
    if dup_mask.any():
        compare_cols = [c for c in raw_combined.columns if c != ID_COL]
        n_conflicting = 0
        conflicting_examples = []
        for idv, group in raw_combined.loc[dup_mask].groupby(ID_COL):
            if group[compare_cols].drop_duplicates().shape[0] > 1:
                n_conflicting += 1
                if len(conflicting_examples) < 5:
                    conflicting_examples.append(idv)
        if n_conflicting > 0:
            log.add(f"WARNING: {n_conflicting} valid but duplicated ids have rows with DISCORDANT VALUES "
                     f"(not just identical copies); the first row found is kept, others discarded. "
                     f"Examples: {conflicting_examples}. Check whether a different tie-break rule is needed.")

    raw_combined = raw_combined.drop_duplicates(subset=ID_COL, keep="first")
    n_after = raw_combined.shape[0]
    if n_before != n_after:
        log.add(f"Merged componenti_ambientali_1 and _2 (after discarding implausible ids): {n_before} rows -> "
                 f"{n_after} rows after de-duplication on id ({n_before - n_after} duplicates removed)")
    else:
        log.add(f"Merged componenti_ambientali_1 and _2 (after discarding implausible ids): {n_after} rows, "
                 f"no duplicate id found")

    full = merge_side_file(raw_combined, fumo, "fumo", "combined dataset", log)
    full = merge_side_file(full, scolarita, "scolarita", "combined dataset", log)

    # ---------------- COHORT ASSIGNMENT FROM GENOTYPE FILES ----------------
    gen1_ids = load_ids_from_genotype_file(args.gen1_ids_file, log, "gen1 (Cohort 1, RES)")
    gen2_ids = load_ids_from_genotype_file(args.gen2_ids_file, log, "gen2 (Cohort 2, ACH)")

    overlap = gen1_ids & gen2_ids
    if overlap:
        log.add(f"WARNING: {len(overlap)} ids present in BOTH genotype files; "
                 f"assigned to Cohort 1 by priority.")

    def assign_cohort(id_val):
        doubled = f"{id_val}_{id_val}"
        if id_val in gen1_ids or doubled in gen1_ids:
            return GROUP1_LABEL
        if id_val in gen2_ids or doubled in gen2_ids:
            return GROUP2_LABEL
        return None

    full[GROUP_COL] = full[ID_COL].astype(str).map(assign_cohort)

    n_unassigned = full[GROUP_COL].isna().sum()
    if n_unassigned > 0:
        log.add(f"WARNING: {n_unassigned} ids in componenti_ambientali not found in EITHER of the two "
                 f"genotype files; they will be excluded from the analysis. Check the id format "
                 f"(expected identical to the genotype one, e.g. 'RES02977_RES02977').")
        full = full.dropna(subset=[GROUP_COL])

    log.add(f"Final dataset: {full.shape[0]} rows, {full.shape[1]} columns. "
            f"Distribution of '{GROUP_COL}': {full[GROUP_COL].value_counts().to_dict()}")

    full = clean_code_like_values(full, COLUMNS_TO_CLEAN_CODE_LIKE, log)

    buffer_vars = detect_buffer_vars(full, seminativi_buffers=seminativi_buffers)
    continuous_vars = [v for v in BASE_CONTINUOUS_VARS if v in full.columns]
    if "education_years" in full.columns:
        continuous_vars.append("education_years")
    continuous_vars += buffer_vars

    categorical_vars = [v for v in BASE_CATEGORICAL_VARS if v in full.columns and v != GROUP_COL]
    if "fumo_boolean" in full.columns:
        categorical_vars.append("fumo_boolean")
    if "education_level" in full.columns:
        categorical_vars.append("education_level")

    log.add(f"Continuous variables analyzed ({len(continuous_vars)}): {continuous_vars}")
    log.add(f"Categorical variables analyzed ({len(categorical_vars)}): {categorical_vars}")

    rows, g1, g2 = build_table1(full, GROUP_COL, continuous_vars, categorical_vars, log)

    csv_path = os.path.join(out_dir, "table1.csv")
    save_table1_csv(rows, g1, g2, csv_path)
    log.add(f"Saved {csv_path}")

    full.to_csv(os.path.join(out_dir, "combined_dataset.csv"), index=False, encoding="utf-8-sig")
    log.add("Saved combined_dataset.csv (full join, useful for downstream analyses)")

    # ---------------- FIGURES ----------------
    appendix_images = []

    for var in [v for v in BASE_CONTINUOUS_VARS + (["education_years"] if "education_years" in full.columns else []) if v in full.columns]:
        p = plot_continuous_box(full, var, GROUP_COL, g1, g2, os.path.join(img_dir, f"boxplot_{var}.png"), log)
        if p:
            appendix_images.append((p, f"Distribution of {var} by cohort"))

    for var in categorical_vars:
        p = plot_categorical_bar(full, var, GROUP_COL, g1, g2, img_dir, log)
        if p:
            appendix_images.append((p, f"Distribution of {var} by cohort"))

    for category in ["seminativi", "vigneti", "risaie"]:
        if category == "seminativi" and seminativi_buffers is not None:
            cat_buffer_cols = [f"seminativi_{b}" for b in seminativi_buffers if f"seminativi_{b}" in full.columns]
        else:
            cat_buffer_cols = [c for c in full.columns if c.startswith(category + "_")]

        if len(cat_buffer_cols) <= 1:
            # Single buffer selected: a grouped/trend-by-buffer plot doesn't
            # make sense, so fall back to a simple single boxplot instead.
            if cat_buffer_cols:
                var = cat_buffer_cols[0]
                p = plot_continuous_box(full, var, GROUP_COL, g1, g2, os.path.join(img_dir, f"boxplot_{var}.png"), log)
                if p:
                    appendix_images.append((p, f"Distribution of {var} by cohort"))
        else:
            p1 = plot_buffer_grouped_box(full, category, GROUP_COL, g1, g2, img_dir, log, buffer_cols=cat_buffer_cols)
            if p1:
                appendix_images.append((p1, f"{category}: boxplot by buffer distance, {g1} vs {g2}"))
            p2 = plot_buffer_pvalue_trend(full, category, GROUP_COL, g1, g2, img_dir, log, buffer_cols=cat_buffer_cols)
            if p2:
                appendix_images.append((p2, f"{category}: p-value trend across buffer distances"))

    log.add(f"Generated {len(appendix_images)} images in {img_dir}")

    # ---------------- DOCX ----------------
    docx_path = os.path.join(out_dir, "table1.docx")
    save_table1_docx(rows, g1, g2, docx_path, images_for_appendix=appendix_images)
    log.add(f"Saved {docx_path}")

    log.save(os.path.join(out_dir, "run_log.txt"))
    log.add("Done.")


if __name__ == "__main__":
    main()